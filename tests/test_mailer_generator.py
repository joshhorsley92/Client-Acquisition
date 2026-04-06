import os
import io
import pytest
from unittest.mock import MagicMock, patch
from docx import Document

@pytest.fixture
def sample_lead():
    return {"id": "uuid-123", "business_name": "Great Shop One", "platform_source": "etsy", "review_count": 250}

@pytest.fixture
def sample_contact():
    return {"name": "Jane Smith", "role": "Owner", "email": "jane@greatshop.com"}

@pytest.fixture
def sample_signals():
    return {"has_website": True, "has_social_media": False, "website_quality": "basic", "has_seo": False, "has_paid_ads": False}

def test_generates_docx_file(sample_lead, sample_contact, sample_signals):
    from outreach.generator import OutreachGenerator
    gen = OutreachGenerator(db=MagicMock(), base_url="https://turnkey.com/start")
    result = gen.generate_mailer(sample_lead, sample_contact, sample_signals)
    assert isinstance(result, bytes)
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0

def test_contains_business_name(sample_lead, sample_contact, sample_signals):
    from outreach.generator import OutreachGenerator
    gen = OutreachGenerator(db=MagicMock(), base_url="https://turnkey.com/start")
    result = gen.generate_mailer(sample_lead, sample_contact, sample_signals)
    doc = Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Great Shop One" in full_text

def test_contains_contact_name(sample_lead, sample_contact, sample_signals):
    from outreach.generator import OutreachGenerator
    gen = OutreachGenerator(db=MagicMock(), base_url="https://turnkey.com/start")
    result = gen.generate_mailer(sample_lead, sample_contact, sample_signals)
    doc = Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane" in full_text

def test_contains_pitch_paragraphs(sample_lead, sample_contact, sample_signals):
    from outreach.generator import OutreachGenerator
    gen = OutreachGenerator(db=MagicMock(), base_url="https://turnkey.com/start")
    result = gen.generate_mailer(sample_lead, sample_contact, sample_signals)
    doc = Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()
    assert "competitors" in full_text or "social" in full_text

def test_contains_brand_elements(sample_lead, sample_contact, sample_signals):
    from outreach.generator import OutreachGenerator
    gen = OutreachGenerator(db=MagicMock(), base_url="https://turnkey.com/start")
    result = gen.generate_mailer(sample_lead, sample_contact, sample_signals)
    doc = Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "YOUR GROWTH, UNLOCKED" in full_text

def test_handles_missing_contact_name(sample_lead, sample_signals):
    from outreach.generator import OutreachGenerator
    gen = OutreachGenerator(db=MagicMock(), base_url="https://turnkey.com/start")
    result = gen.generate_mailer(sample_lead, None, sample_signals)
    doc = Document(io.BytesIO(result))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Business Owner" in full_text
