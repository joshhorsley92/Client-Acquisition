import io
import os
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template
from config import BRAND
from outreach.pitch import generate_pitch_points
from outreach.qr_generator import QRGenerator

logger = logging.getLogger(__name__)


class OutreachGenerator:
    def __init__(self, db, base_url: str):
        self.db = db
        self.qr_gen = QRGenerator(base_url)

    def generate_mailer(self, lead: dict, contact: dict | None, signals: dict) -> bytes:
        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = BRAND["font"]
        font.size = Pt(11)

        # Header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run("TURNKEY MARKETING")
        header_run.bold = True
        header_run.font.size = Pt(20)
        header_run.font.color.rgb = RGBColor.from_string(BRAND["text_color"].lstrip("#"))

        tagline_para = doc.add_paragraph()
        tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline_run = tagline_para.add_run(BRAND["tagline"])
        tagline_run.font.size = Pt(10)
        tagline_run.font.color.rgb = RGBColor.from_string(BRAND["subtle_color"].lstrip("#"))

        doc.add_paragraph("")

        # Greeting
        contact_name = contact["name"] if contact else "Business Owner"
        doc.add_paragraph(f"Dear {contact_name},")

        # Platform-specific opening
        business_name = lead["business_name"]
        review_count = lead.get("review_count", "")
        source = lead.get("platform_source", "")

        if source == "etsy":
            opening = (f"We came across {business_name} on Etsy and were impressed by what you've built — "
                       f"{review_count} reviews speak to real customer satisfaction.")
        elif source == "kickstarter":
            opening = (f"We noticed {business_name} on Kickstarter and were impressed by the traction you've gained — "
                       f"{review_count} backers is a strong signal that your product resonates.")
        else:
            opening = (f"We came across {business_name} in our research of growing businesses in the Metro Detroit area "
                       f"and were impressed by what you're building.")

        doc.add_paragraph(opening)

        # Pitch paragraphs
        pitches = generate_pitch_points(signals)
        for pitch in pitches:
            doc.add_paragraph(pitch)

        # CTA
        doc.add_paragraph(
            "At Turnkey Marketing, we specialize in helping businesses like yours grow their reach "
            "and revenue. We'd love to show you what's possible — no obligation, just a conversation."
        )
        doc.add_paragraph("Scan the QR code below to book a free 15-minute consultation, or reply to this letter directly.")

        # QR Code
        utm_url = self.qr_gen.generate_utm_url(lead["id"], source="mailer")
        qr_image = self.qr_gen.generate_qr_image(utm_url)
        qr_buffer = io.BytesIO()
        qr_image.save(qr_buffer, format="PNG")
        qr_buffer.seek(0)

        qr_para = doc.add_paragraph()
        qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = qr_para.add_run()
        run.add_picture(qr_buffer, width=Inches(1.5))

        # Sign-off
        doc.add_paragraph("")
        doc.add_paragraph("Looking forward to connecting,")
        sign = doc.add_paragraph("The Turnkey Marketing Team")
        sign.runs[0].bold = True

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_email(self, lead: dict, contact: dict | None, signals: dict) -> str:
        contact_name = contact["name"] if contact else "Business Owner"
        business_name = lead["business_name"]
        review_count = lead.get("review_count", "")
        source = lead.get("platform_source", "")
        utm_url = self.qr_gen.generate_utm_url(lead["id"], source="email")

        pitches = generate_pitch_points(signals)

        if source == "etsy":
            opening = f"We came across {business_name} on Etsy — {review_count} reviews is impressive."
        elif source == "kickstarter":
            opening = f"We noticed {business_name} on Kickstarter — {review_count} backers shows real traction."
        else:
            opening = f"We came across {business_name} in our research of growing Metro Detroit businesses."

        pitch_html = "\n".join(f"<p style='color: #333; font-size: 14px; line-height: 1.6;'>{p}</p>" for p in pitches)

        template = Template("""<!DOCTYPE html>
<html>
<body style="margin: 0; padding: 0; background-color: #f5f5f5; font-family: Arial, sans-serif;">
<table width="100%" cellpadding="0" cellspacing="0" style="max-width: 600px; margin: 0 auto; background-color: #ffffff;">
  <tr>
    <td style="padding: 30px 40px; text-align: center; background-color: {{ dark_bg }};">
      <h1 style="color: #ffffff; margin: 0; font-size: 24px;">TURN<span style="color: {{ accent }};">KEY</span></h1>
      <p style="color: {{ subtle }}; font-size: 10px; letter-spacing: 3px; margin: 5px 0 0;">{{ tagline }}</p>
    </td>
  </tr>
  <tr>
    <td style="padding: 30px 40px;">
      <p style="color: #333; font-size: 14px;">Dear {{ contact_name }},</p>
      <p style="color: #333; font-size: 14px; line-height: 1.6;">{{ opening }}</p>
      {{ pitch_html }}
      <p style="color: #333; font-size: 14px; line-height: 1.6;">
        We'd love to show you what's possible — no obligation, just a conversation.
      </p>
      <p style="text-align: center; margin: 30px 0;">
        <a href="{{ utm_url }}" style="background-color: {{ accent }}; color: {{ dark_bg }}; padding: 12px 30px; text-decoration: none; border-radius: 4px; font-weight: bold;">Book a Free Consultation</a>
      </p>
      <p style="color: #333; font-size: 14px;">Looking forward to connecting,</p>
      <p style="color: #333; font-size: 14px;"><strong>The Turnkey Marketing Team</strong></p>
    </td>
  </tr>
</table>
</body>
</html>""")

        return template.render(
            contact_name=contact_name,
            opening=opening,
            pitch_html=pitch_html,
            utm_url=utm_url,
            accent=BRAND["accent_color"],
            dark_bg=BRAND["dark_bg"],
            subtle=BRAND["subtle_color"],
            tagline=BRAND["tagline"],
        )

    def generate_for_lead(self, lead_id: str, format: str = "both") -> dict:
        lead = self.db.get_lead_by_id(lead_id)
        contacts = self.db.get_contacts_for_lead(lead_id)
        signals = self.db.get_signals_for_lead(lead_id) or {}
        contact = contacts[0] if contacts else None

        results = {}
        output_dir = os.path.join("output", lead_id)
        os.makedirs(output_dir, exist_ok=True)

        if format in ("both", "mailer"):
            mailer_bytes = self.generate_mailer(lead, contact, signals)
            mailer_path = os.path.join(output_dir, "mailer.docx")
            with open(mailer_path, "wb") as f:
                f.write(mailer_bytes)
            results["mailer"] = mailer_path

            utm_url = self.qr_gen.generate_utm_url(lead_id, source="mailer")
            self.db.log_outreach({
                "lead_id": lead_id,
                "type": "mailer",
                "utm_code": f"acq_{lead_id[:8]}",
                "qr_url": utm_url,
                "personalization_notes": ", ".join(generate_pitch_points(signals)[:2]),
            })

        if format in ("both", "email"):
            email_html = self.generate_email(lead, contact, signals)
            email_path = os.path.join(output_dir, "email.html")
            with open(email_path, "w") as f:
                f.write(email_html)
            results["email"] = email_path

            utm_url = self.qr_gen.generate_utm_url(lead_id, source="email")
            self.db.log_outreach({
                "lead_id": lead_id,
                "type": "email",
                "utm_code": f"acq_{lead_id[:8]}",
                "qr_url": utm_url,
                "personalization_notes": ", ".join(generate_pitch_points(signals)[:2]),
            })

        logger.info(f"Generated {format} for lead {lead_id}")
        return results
